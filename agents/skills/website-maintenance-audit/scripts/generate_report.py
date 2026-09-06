#!/usr/bin/env python3
import argparse, json, shutil
from pathlib import Path
from docx import Document

SYMBOL={"pass":"✅","warning":"⚠️","fail":"❌","manual":"⚠️","not_verifiable":"—"}
TOKENS={
"navigation":"{{STATUS_NAVIGATION}}","links":"{{STATUS_LINKS}}","impressum":"{{STATUS_IMPRESSUM}}",
"datenschutz":"{{STATUS_DATENSCHUTZ}}","google_fonts":"{{STATUS_GOOGLE_FONTS}}","desktop":"{{STATUS_DESKTOP}}",
"mobile":"{{STATUS_MOBILE}}","tablet_landscape":"{{STATUS_TABLET_LANDSCAPE}}","tablet_portrait":"{{STATUS_TABLET_PORTRAIT}}",
"dsgvo_check":"{{STATUS_DSGVO_CHECK}}","backup":"{{STATUS_BACKUP}}","ssl":"{{STATUS_SSL}}"}

def replace_in_paragraph(p, mapping):
    text=p.text
    changed=False
    for k,v in mapping.items():
        if k in text: text=text.replace(k,str(v)); changed=True
    if changed: p.text=text

def all_paragraphs(doc):
    for p in doc.paragraphs: yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: yield p
    for s in doc.sections:
        for hf in (s.header,s.footer):
            for p in hf.paragraphs: yield p
            for t in hf.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs: yield p

def main():
    ap=argparse.ArgumentParser(); ap.add_argument("results"); ap.add_argument("template"); ap.add_argument("output")
    a=ap.parse_args(); data=json.loads(Path(a.results).read_text(encoding="utf-8")); shutil.copy2(a.template,a.output); doc=Document(a.output)
    byid={x["id"]:x for x in data["results"]}
    wp=byid.get("wordpress",{}); php=byid.get("php",{})
    wpver="—"
    if "Öffentlich erkannte WordPress-Version:" in wp.get("summary",""): wpver=wp["summary"].split(":",1)[1].strip()
    mapping={"{{AUDIT_DATE}}":data["audit_date"],"{{AUDIT_MONTH}}":data["audit_month"],"{{WEBSITE_URL}}":data["site"]["url"],
             "{{PHP_VERSION}}":"—","{{PHP_STATUS}}":"Nicht verifizierbar" if php.get("status") == "not_verifiable" else php.get("summary","—"),"{{WORDPRESS_VERSION}}":wpver,"{{WORDPRESS_STATUS}}":"Nicht verifizierbar" if wp.get("status") == "not_verifiable" else wp.get("summary","—")}
    for id_,token in TOKENS.items(): mapping[token]=SYMBOL.get(byid.get(id_,{}).get("status"),"—")
    findings=[]
    visual_manual=[]
    for x in data["results"]:
        if x["id"] in ("desktop","mobile","tablet_landscape","tablet_portrait") and x["status"] in ("manual","warning","not_verifiable"):
            visual_manual.append(x["label"]); continue
        if x["status"] in ("warning","fail"):
            findings.append(f"{SYMBOL.get(x['status'],'—')} {x['label']}: {x['summary']}")
        elif x["status"] in ("manual","not_verifiable") and x["id"] in ("php","wordpress","backup","dsgvo_check","google_fonts"):
            findings.append(f"{SYMBOL.get(x['status'],'—')} {x['label']}: {x['summary']}")
    if visual_manual:
        findings.append("⚠️ Responsive-Darstellung: visuelle Browserprüfung offen für " + ", ".join(visual_manual) + ".")
    mapping["{{FINDINGS}}"]="\n".join(findings[:7]) if findings else "Keine Auffälligkeiten."
    for p in all_paragraphs(doc): replace_in_paragraph(p,mapping)
    doc.save(a.output); print(a.output)
if __name__=="__main__": main()
